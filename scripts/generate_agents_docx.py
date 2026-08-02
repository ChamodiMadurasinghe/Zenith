"""Generate ChequeMate / Zenith agents reference as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "ChequeMate-Agents-Reference.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_rule_row(table, rule: str, detail: str) -> None:
    row = table.add_row()
    row.cells[0].text = rule
    row.cells[1].text = detail


def agent_section_break(doc: Document) -> None:
    """Visual separator between agent sections."""
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()

    # --- Document title ---
    title = doc.add_heading("ChequeMate Agents Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("What each agent does and which rules apply")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph(
        "This document describes every AI agent in the Zenith / ChequeMate project: "
        "the four-agent invoice orchestration pipeline (agentic/) and the legacy web-app agents (agents/). "
        "Each agent has its own section below."
    )

    add_heading(doc, "Two Agent Naming Schemes", 1)
    doc.add_paragraph(
        "The project uses two overlapping numbering systems. The agentic invoice pipeline (POST /api/orchestrate) "
        "labels agents 1–4 for vision, anomaly, liquidity, and dealer liaison. The main web app README labels "
        "Agent 1 = ingestion, Agent 2 = bundling assistant, Agent 3 = analyst. This document covers both."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Context"
    hdr[1].text = "Agent 1"
    hdr[2].text = "Agent 2"
    row = table.add_row()
    row.cells[0].text = "Agentic pipeline (invoice workflow)"
    row.cells[1].text = "Vision extraction"
    row.cells[2].text = "Anomaly audit"
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Agent 3 = Liquidity forecast"
    row.cells[2].text = "Agent 4 = Dealer liaison"
    row = table.add_row()
    row.cells[0].text = "Legacy web app (README)"
    row.cells[1].text = "Ingestion (Gemini vision)"
    row.cells[2].text = "Cheque Assistant (bundling chat)"
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = "Agent 3 = Analyst (reports)"
    row.cells[2].text = "Plus: Guide, Reviewer"

    # ========================================================================
    # AGENTIC PIPELINE — AGENT 1
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 1 — Vision Extractor", 0)
    add_heading(doc, "Agentic Invoice Pipeline", 1)

    doc.add_paragraph(
        "Reads a supplier invoice image and produces a structured InvoiceDraft "
        "(supplier name, total LKR, due date, line items, dealer ID)."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Triggered when the orchestrator receives an INVOICE_IMAGE event.")
    add_bullet(doc, "Calls Gemini Vision via agents/ingestion.py → extract_invoice().")
    add_bullet(doc, "Maps extracted JSON into an InvoiceDraft used by Agents 2 and 3.")
    add_bullet(doc, "Runs inside the PER loop (Plan → Execute → Review → Decide).")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "Tool: ZenithVisionExtractor in agentic/adapters/zenith_tools.py", "Code: ")
    add_bullet(doc, "Pipeline: InvoicePipeline._run_agent1() in agentic/orchestrator/pipeline.py", "Code: ")
    add_bullet(doc, "LLM: Gemini Vision — always used (not optional)", "AI: ")

    add_heading(doc, "Rules That Apply", 2)
    rules = doc.add_table(rows=1, cols=2)
    rules.style = "Table Grid"
    rules.rows[0].cells[0].text = "Rule"
    rules.rows[0].cells[1].text = "Detail"
    add_rule_row(rules, "PER review — supplier_name", "Must be non-empty")
    add_rule_row(rules, "PER review — total_lkr", "Must be greater than 0")
    add_rule_row(rules, "Retry limit", "Up to 3 attempts (PERLoop.MAX_RETRIES)")
    add_rule_row(rules, "FSM on retry", "EXTRACTING ↔ RETRYING oscillation")
    add_rule_row(rules, "FSM on success", "Transition to AUDITING")
    add_rule_row(rules, "Extraction prompt", "Extract only visible fields — do not invent line items")
    add_rule_row(rules, "Decision outcomes", "continue | retry | fail")

    add_heading(doc, "Trace Label", 2)
    doc.add_paragraph("agent1 — plan / execute / review / decide")

    # ========================================================================
    # AGENTIC PIPELINE — AGENT 2
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 2 — Smart Anomaly Guard", 0)
    add_heading(doc, "Agentic Invoice Pipeline", 1)

    doc.add_paragraph(
        "Audits the extracted invoice against dealer history and business rules. "
        "Can lock the invoice and stop the entire pipeline before any cheque date is suggested."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Step 1 (rules): Runs check_invoice_anomalies() against SQLite dealer history.")
    add_bullet(doc, "Step 2 (conditional AI): Calls Gemini only when flags or severity warrant review.")
    add_bullet(doc, "On lock: FSM moves to LOCKED (terminal) — Agent 3 never runs.")
    add_bullet(doc, "On pass: FSM moves to FORECASTING.")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "Rules tool: ZenithAnomalyGuard in agentic/adapters/zenith_tools.py", "Code: ")
    add_bullet(doc, "Anomaly logic: agents/anomaly.py", "Code: ")
    add_bullet(doc, "Conditional AI: agentic/adapters/agent_ai.py → ai_audit_review()", "Code: ")
    add_bullet(doc, "LLM: Gemini text — only when AGENT_CONDITIONAL_AI=true and triggers fire", "AI: ")

    add_heading(doc, "Deterministic Rules (agents/anomaly.py)", 2)
    anomaly = doc.add_table(rows=1, cols=3)
    anomaly.style = "Table Grid"
    anomaly.rows[0].cells[0].text = "Check"
    anomaly.rows[0].cells[1].text = "Severity"
    anomaly.rows[0].cells[2].text = "Effect"
    checks = [
        ("Supplier present but amount ≤ 0", "high", "Flag"),
        ("Invoice date > 30 days in future", "high", "Flag"),
        ("Invoice date > 1 year old", "medium", "Flag"),
        ("Duplicate invoice number for dealer", "high", "Flag"),
        ("Amount > 3× dealer historical average", "medium", "Flag"),
    ]
    for check, sev, effect in checks:
        row = anomaly.add_row()
        row.cells[0].text = check
        row.cells[1].text = sev
        row.cells[2].text = effect

    add_heading(doc, "Lock Rule", 2)
    doc.add_paragraph(
        "Any flag with severity == \"high\" sets locked = True. "
        "The pipeline stops and the invoice reaches the LOCKED terminal state."
    )

    add_heading(doc, "When Gemini (agent2-ai) Runs", 2)
    add_bullet(doc, "AGENT_CONDITIONAL_AI=true (default in .env)")
    add_bullet(doc, "Rules result is locked, OR any flags exist, OR severity is warning/critical")
    add_bullet(doc, "On API failure: falls back to rules result (no crash)")

    add_heading(doc, "Trace Labels", 2)
    doc.add_paragraph("agent2 (rules) → optionally agent2-ai (Gemini review)")

    # ========================================================================
    # AGENTIC PIPELINE — AGENT 3
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 3 — Liquidity & Cheque Forecaster", 0)
    add_heading(doc, "Agentic Invoice Pipeline", 1)

    doc.add_paragraph(
        "Computes the optimal cheque date to maximize legal float while meeting supplier deadlines, "
        "using CBSL bank holidays and the liquidity engine."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Step 1 (engine): Scans up to 45 days backward, skips Sundays, applies apply_liquidity_dates().")
    add_bullet(doc, "Picks the date with maximum float days before true settlement.")
    add_bullet(doc, "Step 2 (conditional AI): Gemini picks date and writes SME-facing rationale for strategic cases.")
    add_bullet(doc, "Re-runs when dealer rejects date and provides an alternative (REFORECASTING state).")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "Engine: ZenithLiquidityForecaster in agentic/adapters/zenith_tools.py", "Code: ")
    add_bullet(doc, "Core math: core/liquidity_engine.py", "Code: ")
    add_bullet(doc, "Conditional AI: agentic/adapters/agent_ai.py → ai_forecast_review()", "Code: ")
    add_bullet(doc, "Holidays: loaded from DB via repository.get_holidays()", "Data: ")

    add_heading(doc, "Engine Rules", 2)
    add_bullet(doc, "Respect supplier deadline / due date from Agent 1 draft")
    add_bullet(doc, "Respect dealer alternative pickup date on re-forecast")
    add_bullet(doc, "Maximize float_days (days gained by holiday lag)")
    add_bullet(doc, "Return top 5 candidate dates for AI review when triggered")

    add_heading(doc, "When Gemini (agent3-ai) Runs", 2)
    add_bullet(doc, "AGENT_CONDITIONAL_AI=true")
    add_bullet(doc, "Dealer re-forecast (negotiation_round > 0 or alternative pickup date set)")
    add_bullet(doc, "Invoice amount ≥ 500,000 LKR")
    add_bullet(doc, "Float ≥ 1 day")
    add_bullet(doc, "More than 1 candidate date")
    add_bullet(doc, "On API failure: falls back to engine plan")

    add_heading(doc, "FSM Transitions", 2)
    add_bullet(doc, "Success → AWAITING_DEALER")
    add_bullet(doc, "Dealer NO + date → REFORECASTING → Agent 3 rerun → AWAITING_DEALER")

    add_heading(doc, "Trace Labels", 2)
    doc.add_paragraph("agent3 (engine) → optionally agent3-ai (Gemini strategic review)")

    # ========================================================================
    # AGENTIC PIPELINE — AGENT 4
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 4 — Dealer Liaison", 0)
    add_heading(doc, "Agentic Invoice Pipeline", 1)

    doc.add_paragraph(
        "Drafts dealer confirmation messages and parses replies. "
        "Coordinates the first human gate — dealer must confirm or reject the proposed cheque date."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Drafts outbound message with recommended date and amount.")
    add_bullet(doc, "Parses dealer reply via templates and regex (no LLM in prototype).")
    add_bullet(doc, "Routes confirmed replies to merchant approval gate.")
    add_bullet(doc, "Routes rejected replies with alternative date back to Agent 3 for re-forecast.")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "ZenithDealerLiaison in agentic/adapters/zenith_tools.py", "Code: ")
    add_bullet(doc, "LLM: None — templates + regex only", "AI: ")

    add_heading(doc, "Outbound Message Format", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Please confirm cheque pickup on {date} for LKR {amount}. '
        'Reply YES to confirm or NO with your preferred date."'
    )
    run.italic = True

    add_heading(doc, "Reply Parsing Rules", 2)
    replies = doc.add_table(rows=1, cols=3)
    replies.style = "Table Grid"
    replies.rows[0].cells[0].text = "Dealer Reply"
    replies.rows[0].cells[1].text = "Parsed Status"
    replies.rows[0].cells[2].text = "FSM Transition"
    reply_rows = [
        ("YES, OK, confirm, confirmed", "confirmed", "→ AWAITING_APPROVAL"),
        ('NO + date (e.g. "NO April 15", ISO date)', "alternative_date", "→ REFORECASTING → Agent 3"),
        ("Unclear / no date", "pending", "Stay in AWAITING_DEALER"),
    ]
    for reply, status, fsm in reply_rows:
        row = replies.add_row()
        row.cells[0].text = reply
        row.cells[1].text = status
        row.cells[2].text = fsm

    add_heading(doc, "Channels", 2)
    add_bullet(doc, "Web simulate buttons → DEALER_REPLY event")
    add_bullet(doc, "WhatsApp text → whatsapp_bridge → same DEALER_REPLY event")

    add_heading(doc, "Trace Label", 2)
    doc.add_paragraph("agent4 — plan / execute (reply handling logs execute only)")

    # ========================================================================
    # UNIVERSAL ORCHESTRATOR RULES
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Universal Orchestrator Rules", 0)
    add_heading(doc, "PER Loop, FSM, and Human Gates", 1)

    add_heading(doc, "PER Loop (Plan → Execute → Review → Decide)", 2)
    doc.add_paragraph("Every agentic agent step runs through agentic/orchestrator/per_loop.py:")
    add_bullet(doc, "Plan — log inputs and strategy")
    add_bullet(doc, "Execute — call tool (vision / rules / engine / liaison)")
    add_bullet(doc, "Review — validate output against business rules")
    add_bullet(doc, "Decide — continue | retry | lock | fail")
    add_bullet(doc, "Max 3 retries on execute exception or failed review")
    add_bullet(doc, "All steps appended to trace — never overwritten")

    add_heading(doc, "Invoice FSM (12 States, 3 Terminal Outcomes)", 2)
    fsm = doc.add_table(rows=1, cols=2)
    fsm.style = "Table Grid"
    fsm.rows[0].cells[0].text = "Terminal State"
    fsm.rows[0].cells[1].text = "Meaning"
    for state, meaning in [
        ("COMPLETED", "Merchant approved — cheque plan is final"),
        ("REJECTED", "Merchant rejected at approval gate"),
        ("LOCKED", "Anomaly guard blocked invoice — no forecast"),
    ]:
        row = fsm.add_row()
        row.cells[0].text = state
        row.cells[1].text = meaning

    doc.add_paragraph(
        "Non-terminal flow: RECEIVED → EXTRACTING → AUDITING → FORECASTING → "
        "AWAITING_DEALER → AWAITING_APPROVAL → COMPLETED. "
        "Retry path: EXTRACTING ↔ RETRYING. Re-forecast path: AWAITING_DEALER → REFORECASTING."
    )

    add_heading(doc, "Human Gates (Orchestrator — Not Agents)", 2)
    gates = doc.add_table(rows=1, cols=3)
    gates.style = "Table Grid"
    gates.rows[0].cells[0].text = "Gate"
    gates.rows[0].cells[1].text = "Event"
    gates.rows[0].cells[2].text = "Who Decides"
    for gate, event, who in [
        ("Gate 1 — Dealer", "DEALER_REPLY", "Dealer confirms or rejects date"),
        ("Gate 2 — Merchant", "APPROVAL_DECISION { approved: true/false }", "Merchant approves or rejects"),
    ]:
        row = gates.add_row()
        row.cells[0].text = gate
        row.cells[1].text = event
        row.cells[2].text = who

    add_heading(doc, "Configuration Flags", 2)
    add_bullet(doc, "USE_AGENTIC_ORCHESTRATOR=true — routes WhatsApp/web to agentic pipeline", "Env: ")
    add_bullet(doc, "AGENT_CONDITIONAL_AI=false — disables Gemini escalation for Agents 2 and 3", "Env: ")

    add_heading(doc, "API Entry Points", 2)
    add_bullet(doc, "POST /api/orchestrate — send InboundEvent JSON")
    add_bullet(doc, "GET /api/sessions/{id}/trace — poll agent activity trace")

    # ========================================================================
    # LEGACY — AGENT 1 INGESTION
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 1 — Ingestion", 0)
    add_heading(doc, "Legacy Web App", 1)

    doc.add_paragraph(
        "Gemini vision extraction for the main web upload and WhatsApp intake flows. "
        "Same underlying tool used by the agentic Agent 1."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Extracts invoice/cheque fields from uploaded images.")
    add_bullet(doc, "Returns structured JSON for human verification on the review screen.")
    add_bullet(doc, "Used by routes/ingestion.py and core/whatsapp_intake.py.")

    add_heading(doc, "Rules That Apply", 2)
    add_bullet(doc, "Extract only information visible on the image")
    add_bullet(doc, "Do not invent line items")
    add_bullet(doc, "Support both supplier invoices and handwritten/printed cheques")
    add_bullet(doc, "Return JSON: invoice_no, supplier_name, total_amount, line_items, credit_period_days, etc.")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "agents/ingestion.py → extract_invoice()", "Code: ")
    add_bullet(doc, "Gemini Vision via agents/base.py → generate_with_image()", "AI: ")

    # ========================================================================
    # LEGACY — AGENT 2 ASSISTANT
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 2 — Cheque Assistant", 0)
    add_heading(doc, "Legacy Web App", 1)

    doc.add_paragraph(
        "Conversational cheque bundling assistant on the Cheques page. "
        "Helps merchants divide, group, and date cheques under an LKR ceiling."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Explains Python-computed bundle groups conversationally.")
    add_bullet(doc, "Emits structured proposed_actions JSON for Python to apply.")
    add_bullet(doc, "Supports English, Sinhala, and Tamil.")
    add_bullet(doc, "Pre-fills dealer setup forms for unknown suppliers.")

    add_heading(doc, "Rules That Apply", 2)
    add_bullet(doc, "Must include proposed_actions JSON block when user asks to bundle/divide/split/group")
    add_bullet(doc, "Never say \"hold on\" or \"let me calculate\" without emitting JSON")
    add_bullet(doc, "Ground answers in full bundling context (invoices, ceiling, groups, dates)")
    add_bullet(doc, "Python applies and verifies actions — agent never writes to DB directly")
    add_bullet(doc, "List verification issues clearly; explain risks but do not refuse preview")

    add_heading(doc, "Allowed Actions (proposed_actions)", 2)
    actions = [
        "divide_into_cheques — split invoices into N cheques",
        "create_bundles — assign invoice groups with cheque dates",
        "assign_invoices — map invoice IDs to group numbers",
        "set_cheque_date — set date for a cheque group",
        "move_invoice — move invoice between groups",
        "postpone_cheque — delay a cheque group by N days",
        "split_invoice — split a single invoice",
        "recalculate_dates — recompute dates from current state",
    ]
    for action in actions:
        add_bullet(doc, action)

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "agents/assistant.py", "Code: ")
    add_bullet(doc, "OpenAI gpt-3.5-turbo (or configured OPENAI_CHAT_MODEL)", "AI: ")

    # ========================================================================
    # LEGACY — AGENT 3 ANALYST
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Agent 3 — Analyst", 0)
    add_heading(doc, "Legacy Web App", 1)

    doc.add_paragraph(
        "Generates markdown financial reports from pre-computed business metrics "
        "after cheque commits. Runs asynchronously."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Receives structured metrics from Python SQL aggregation.")
    add_bullet(doc, "Writes markdown covering cash position, bank trends, inventory velocity.")
    add_bullet(doc, "Report cached in analyst_reports table and shown on Analytics page.")

    add_heading(doc, "Rules That Apply", 2)
    add_bullet(doc, "Use headings, bullet points, and LKR amounts formatted with commas")
    add_bullet(doc, "Cover: cash position vs liabilities, bank deposit trends, invoice velocity, recommendations")
    add_bullet(doc, "Python owns all metric computation — agent only narrates")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "agents/analyst.py → generate_report()", "Code: ")
    add_bullet(doc, "OpenAI (OPENAI_ANALYST_MODEL)", "AI: ")

    # ========================================================================
    # LEGACY — GUIDE
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Guide — App Navigation Helper", 0)
    add_heading(doc, "Legacy Web App", 1)

    doc.add_paragraph(
        "Patient user-manual and technical helper for merchants using the Zenith web app. "
        "Explains screens, buttons, and troubleshooting — not financial decisions."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Navigates users through Invoices, Cheques, Bank Balance, Reports")
    add_bullet(doc, "Explains what each screen and button does")
    add_bullet(doc, "Troubleshoots upload quality, login, language, pending invoices")
    add_bullet(doc, "May perform safe app actions: navigate to section, logout")

    add_heading(doc, "Rules That Apply", 2)
    add_bullet(doc, "Must NOT bundle cheques, split invoices, set dates, or move invoices")
    add_bullet(doc, "Must NOT use proposed_actions or any bundling commands")
    add_bullet(doc, "Direct bundling requests to Agent 2 (Cheque Assistant) on the Cheques page")
    add_bullet(doc, "guide_actions limited to: navigate (with target) and logout")
    add_bullet(doc, "On Cheques page: do not emit guide_actions — user is already there")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "agents/guide.py", "Code: ")
    add_bullet(doc, "OpenAI (configured chat model)", "AI: ")

    # ========================================================================
    # LEGACY — REVIEWER
    # ========================================================================
    agent_section_break(doc)

    add_heading(doc, "Reviewer — Liquidity Review Agent", 0)
    add_heading(doc, "Legacy Web App", 1)

    doc.add_paragraph(
        "SME-style liquidity reviewer that evaluates Python-proposed cheque bundles "
        "from the merchant's perspective — maximizing legal cheque float."
    )

    add_heading(doc, "What It Does", 2)
    add_bullet(doc, "Reviews proposed bundles using full context: dealer profile, holidays, liquidity metrics.")
    add_bullet(doc, "Issues VERDICT: approve or VERDICT: suggest_changes.")
    add_bullet(doc, "In apply mode: emits proposed_actions to implement review suggestions.")

    add_heading(doc, "Rules That Apply", 2)
    add_bullet(doc, "Primary objective: maximize legal cheque float")
    add_bullet(doc, "Cite specific invoice numbers, amounts, groups, and dates from context only")
    add_bullet(doc, "Never invent data not in the context")
    add_bullet(doc, "Start reply with exactly one line: VERDICT: approve OR VERDICT: suggest_changes")
    add_bullet(doc, "Review mode: text suggestions only, no JSON actions")
    add_bullet(doc, "Apply mode: output proposed_actions JSON + summary")

    add_heading(doc, "Evaluation Criteria", 2)
    add_bullet(doc, "stated cheque date vs true_settlement_date vs target_funding_date")
    add_bullet(doc, "days_gained_by_holiday_lag")
    add_bullet(doc, "is_interbank (+1 business day when banks differ)")
    add_bullet(doc, "dealer casual_days and impossible_days")
    add_bullet(doc, "LKR ceiling forcing earlier funding than necessary")

    add_heading(doc, "Implementation", 2)
    add_bullet(doc, "agents/reviewer.py", "Code: ")
    add_bullet(doc, "OpenAI (configured chat model)", "AI: ")

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Generated for ChequeMate / Zenith. Source of truth: agentic/state/invoice_fsm.py, "
        "agentic/orchestrator/pipeline.py, agents/*. See agentic/README.md for API details."
    )
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return doc


def main() -> None:
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
