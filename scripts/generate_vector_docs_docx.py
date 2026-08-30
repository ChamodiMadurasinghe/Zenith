"""Generate Zenith Vector Pattern Engine reference as a Word document.

Matches docs/Zenith-Database-Reference.docx styling.
Diagrams: embedded images + browser links (mermaid.ink — any browser, no account).
"""

from __future__ import annotations

import base64
import json
import re
import zlib
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "Vector_Pattern_Engine.md"
DIAGRAMS_DIR = ROOT / "docs" / "diagrams" / "vector"
OUTPUT = ROOT / "docs" / "Zenith-Vector-Pattern-Engine-Reference.docx"

GITHUB_RAW_BASE = (
    "https://raw.githubusercontent.com/ChamodiMadurasinghe/Zenith/main/docs/diagrams/vector"
)

MERMAID_INK_TIMEOUT = 45

MERMAID_DIAGRAMS: list[tuple[str, str, str]] = [
    (
        "01-system-overview",
        "High-level domain map",
        """flowchart TB
    subgraph sqlite [SQLite source of truth]
        INV[invoices]
        CHQ[cheque]
        ALLOC[cheque_invoice_allocation]
    end
    subgraph python [Python pattern builder]
        DP[dealer_patterns.py]
        VS[vector_store.py]
    end
    subgraph store [Vector store]
        CHROMA[(ChromaDB database/chroma)]
    end
    subgraph ai [Bundling Assistant]
        TOOL[get_dealer_historical_payment_patterns]
    end
    sqlite --> DP
    DP --> VS
    VS --> CHROMA
    CHROMA --> TOOL
    COMMIT[Cheque commit] --> sqlite
    COMMIT --> VS""",
    ),
    (
        "02-data-lifecycle",
        "Business lifecycle",
        """flowchart LR
    S1[Intake] --> S2[Verify]
    S2 --> S3[Draft bundle]
    S3 --> S4[Preview]
    S4 --> S5[User commits]
    S5 --> S6[Vector refresh]
    S6 --> S7[Next chat query]
    S3 -.->|AI reads patterns| PAT[Pattern tool]""",
    ),
    (
        "03-guardrails-boundary",
        "Guardrails boundary",
        """flowchart LR
    VEC[Vector patterns] -->|suggestions only| AI[Bundling Assistant]
    AI -->|dry_run tools| PY[core/bundling.py]
    PY --> GR[collect_bundle_issues]
    GR --> UI[Preview UI]
    UI -->|user Save| SQL[SQLite commit]
    SQL --> REF[Vector refresh]""",
    ),
]


def mermaid_pako(code: str) -> str:
    state = {
        "code": code.strip() + "\n",
        "mermaid": {"theme": "default"},
        "autoSync": True,
        "updateDiagram": True,
    }
    raw = json.dumps(state, separators=(",", ":")).encode("utf-8")
    compressed = zlib.compress(raw, 9)
    return base64.urlsafe_b64encode(compressed).decode("ascii").rstrip("=")


def mermaid_ink_image_url(code: str) -> str:
    return f"https://mermaid.ink/img/pako:{mermaid_pako(code)}"


def github_raw_image_url(stem: str, suffix: str) -> str:
    return f"{GITHUB_RAW_BASE}/{stem}{suffix}"


def fetch_and_save_image(code: str, dest_stem: Path) -> Path | None:
    url = mermaid_ink_image_url(code)
    try:
        resp = requests.get(url, timeout=MERMAID_INK_TIMEOUT)
    except requests.RequestException:
        return None
    if resp.status_code != 200 or len(resp.content) < 200:
        return None
    data = resp.content
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        path = dest_stem.with_suffix(".png")
    elif data[:3] == b"\xff\xd8\xff":
        path = dest_stem.with_suffix(".jpg")
    else:
        return None
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return path


def find_saved_image(stem: str) -> Path | None:
    for ext in (".jpg", ".png"):
        p = DIAGRAMS_DIR / f"{stem}{ext}"
        if p.is_file():
            return p
    return None


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def fill_header_row(row, headers: list[str]) -> None:
    for i, h in enumerate(headers):
        row.cells[i].text = h
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(row.cells[i], "E8EDF4")


def add_grid(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    fill_header_row(table.rows[0], headers)
    for r_i, values in enumerate(rows):
        for c_i, val in enumerate(values):
            table.rows[r_i + 1].cells[c_i].text = val
    doc.add_paragraph("")


def add_diagram_block(
    doc: Document,
    stem: str,
    heading: str,
    mermaid_code: str,
    image_path: Path | None,
) -> tuple[str, str, str, bool]:
    ink_url = mermaid_ink_image_url(mermaid_code)
    suffix = image_path.suffix if image_path else ".jpg"
    github_url = github_raw_image_url(stem, suffix)
    embedded = image_path is not None and image_path.is_file()

    doc.add_heading(heading, 2)
    if embedded:
        doc.add_picture(str(image_path), width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(heading)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        doc.add_paragraph(
            "Diagram image could not be downloaded. Use the browser link below — "
            "it opens as a picture in Chrome, Edge, Firefox, or Safari."
        )

    p = doc.add_paragraph("Open in any browser (image): ")
    add_hyperlink(p, heading + " — view now", ink_url)
    p2 = doc.add_paragraph("After git push, GitHub copy: ")
    add_hyperlink(p2, github_url, github_url)
    local_name = f"{stem}{suffix}" if image_path else f"{stem}.jpg"
    doc.add_paragraph(
        f"Local file: docs/diagrams/vector/{local_name} — attach to WhatsApp, email, or Slack."
    )
    return heading, ink_url, github_url, embedded


def build_document(
    saved: dict[str, Path | None],
) -> tuple[Document, list[tuple[str, str, str, bool]]]:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Zenith Vector Pattern Engine Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("ChromaDB pattern memory, diagrams, and operations catalog")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True

    doc.add_paragraph(
        "This document describes the Vector Pattern Engine in Zenith / ChequeMate — "
        "semantic memory of how each supplier (dealer) was paid in the past. "
        "SQLite remains the source of truth; ChromaDB stores searchable pattern summaries."
    )
    doc.add_paragraph(
        "Diagrams are embedded below. Blue links open a picture in any web browser — "
        "no draw.io, no login, no special software."
    )

    doc.add_heading("Shareable diagram links", 1)
    doc.add_paragraph(
        "Share the Word file, the image files in docs/diagrams/vector/, or the browser links below."
    )

    diagram_results: list[tuple[str, str, str, bool]] = []
    for stem, heading, code in MERMAID_DIAGRAMS:
        result = add_diagram_block(doc, stem, heading, code, saved.get(stem))
        diagram_results.append(result)

    doc.add_heading("Introduction", 1)
    add_grid(
        doc,
        ["Aspect", "Detail"],
        [
            ["What it is", "One ChromaDB text summary per dealer, refreshed on cheque commit"],
            ["What it is not", "Not source of truth, not auto-commit, not date/ceiling math"],
            ["One-line summary", "SQLite = facts; ChromaDB = searchable pattern summaries for AI"],
        ],
    )

    doc.add_heading("Why a vector database?", 1)
    doc.add_paragraph(
        "Plain SQL lists history; the AI needs natural-language context. ChromaDB + OpenAI "
        "embeddings retrieve a pre-written pattern document. All dates and ceilings still "
        "pass through core/guardrails.py."
    )

    doc.add_heading("Payment classification", 1)
    add_grid(
        doc,
        ["Type", "Rule", "Aging in document"],
        [
            ["Bundled", "Cheque has 2+ distinct invoices", "Average aging across bundled rows"],
            ["Unbundled", "Exactly 1 invoice, single part", "Exact per-invoice lines — no average"],
            ["Split", "Same invoice on 2+ cheques", "Per-part aging + payment pattern line"],
            ["Mixed", "Both bundled and unbundled history", "Both sections appear"],
        ],
    )

    doc.add_heading("Aging rules", 1)
    doc.add_paragraph("Formula: aging_days = clearance_date - invoiced_date")
    doc.add_paragraph(
        "Clearance: predicted_clearance_date → deposit_timetable.target_funding_date → cheque_date"
    )

    doc.add_heading("Example pattern document", 2)
    example = doc.add_paragraph()
    run = example.add_run(
        "Dealer: ABD Traders (ID: 1)\nBundling History: mixed\n\n"
        "Aging Analysis:\n"
        "- Bundled Invoices Average Aging: 38 days (4 multi-invoice cheques).\n"
        "- Unbundled Invoice Records: Inv #101 (21 days), Inv #104 (14 days).\n\n"
        "Preferred Paying Account: Commercial Bank — 12 of 15 cheques (80%).\n"
        "Payment Pattern: Bills over 500k LKR split into 2 parts, 7-day gap."
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_heading("ChromaDB storage model", 1)
    add_grid(
        doc,
        ["Setting", "Value"],
        [
            ["Path", "database/chroma/"],
            ["Collection", "dealer_payment_patterns"],
            ["Document ID", "dealer-{dealer_id}"],
            ["Embeddings", "OpenAI text-embedding-3-small"],
        ],
    )

    doc.add_heading("Configuration (.env)", 1)
    add_grid(
        doc,
        ["Variable", "Default", "Purpose"],
        [
            ["ENABLE_VECTOR_PATTERNS", "true", "Master switch"],
            ["CHROMA_PERSIST_DIR", "database/chroma", "Storage path"],
            ["OPENAI_EMBEDDING_MODEL", "text-embedding-3-small", "Embedding model"],
            ["PATTERN_LARGE_BILL_LKR", "500000", "Split-pattern threshold"],
            ["OPENAI_API_KEY", "—", "Required for embeddings"],
            ["USE_FAKE_AI", "false", "Mock patterns; skips Chroma"],
        ],
    )

    doc.add_heading("Setup and operations", 1)
    setup = doc.add_paragraph()
    run = setup.add_run("pip install -r requirements.txt\npython scripts/backfill_dealer_patterns.py")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    doc.add_paragraph("Auto-refresh on every successful cheque commit.")

    doc.add_heading("Guardrails boundary", 1)
    doc.add_paragraph(
        "Vector DB CANNOT set dates, override ceiling, skip holidays, or commit without user Save."
    )

    doc.add_heading("File reference", 1)
    add_grid(
        doc,
        ["File", "Purpose"],
        [
            ["core/dealer_patterns.py", "Build pattern text"],
            ["core/vector_store.py", "Chroma upsert / query"],
            ["agents/bundling_tools.py", "Historical patterns tool"],
            ["docs/diagrams/vector/*.jpg", "Shareable diagram images"],
        ],
    )

    end = doc.add_paragraph("Regenerate: python scripts/generate_vector_docs_docx.py")
    if end.runs:
        end.runs[0].italic = True

    return doc, diagram_results


def update_markdown_diagram_links(saved: dict[str, Path | None]) -> None:
    if not MD_PATH.is_file():
        return
    lines = [
        "| Diagram | Open in browser (any browser) | GitHub copy (after push) | Local file |",
        "|---------|------------------------------|---------------------------|------------|",
    ]
    for stem, heading, code in MERMAID_DIAGRAMS:
        ink = mermaid_ink_image_url(code)
        img = saved.get(stem)
        suffix = img.suffix if img else ".jpg"
        gh = github_raw_image_url(stem, suffix)
        local = f"docs/diagrams/vector/{stem}{suffix}"
        lines.append(f"| {heading} | [View image]({ink}) | [GitHub]({gh}) | `{local}` |")

    block = "## 3. Architecture diagrams\n\n" + "\n".join(lines) + "\n\n"
    md = MD_PATH.read_text(encoding="utf-8")
    md = re.sub(r"## 3\. Architecture diagrams[\s\S]*?(?=## 4\.)", block, md, count=1)
    md = re.sub(
        r"Offline.*?\n\n> Links work after",
        "Teammates can open the **View image** links in any browser, attach the local `.jpg` files, or read diagrams embedded in the Word doc.\n\n> GitHub links work after",
        md,
        count=1,
    )
    MD_PATH.write_text(md, encoding="utf-8")


def main() -> None:
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    saved: dict[str, Path | None] = {}
    for stem, _heading, code in MERMAID_DIAGRAMS:
        saved[stem] = fetch_and_save_image(code, DIAGRAMS_DIR / stem)

    doc, results = build_document(saved)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    update_markdown_diagram_links(saved)

    print(f"Wrote {OUTPUT}")
    print(f"Saved images under {DIAGRAMS_DIR}")
    print()
    print("Share with teammates — opens in any browser, no special software:")
    for heading, ink_url, gh_url, embedded in results:
        flag = "embedded in Word" if embedded else "link only"
        print(f"  [{flag}] {heading}")
        print(f"    Browser: {ink_url}")
        print(f"    GitHub:  {gh_url}")
        print()


if __name__ == "__main__":
    main()
