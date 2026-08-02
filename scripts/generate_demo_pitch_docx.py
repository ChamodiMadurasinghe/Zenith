"""Generate Idealize demo presenter + submission Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Zenith_Video_Pitch_Presenter_Guide.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_quote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True


def add_action(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("On screen: ")
    run.bold = True
    p.add_run(text)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Zenith — Idealize Demo Video Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Presenter cue sheet + submission checklist  |  Max length: 5 minutes  |  Language: English"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True

    meta = doc.add_paragraph()
    meta.add_run("Suggested video file name: ").bold = True
    meta.add_run("Zenith_Video_Pitch")
    meta2 = doc.add_paragraph()
    meta2.add_run("Recording: ").bold = True
    meta2.add_run(
        "Zoom screen share + camera picture-in-picture (presenter visible and audible). "
        "Casual attire allowed. Open Category — show AI agent inputs, steps, and outputs."
    )

    add_heading(doc, "1. Submission requirements (checklist)", 1)
    add_bullet(doc, "Screen recording demonstrating and explaining the application (≤ 5 minutes).")
    add_bullet(doc, "Navigate through each major page/feature with voiceover.")
    add_bullet(doc, "Open Category: show the AI Agent in action — inputs, steps, outputs.")
    add_bullet(doc, "Presenter clearly audible and visible (e.g. Zoom PiP).")
    add_bullet(doc, "Language of use: English.")
    add_bullet(
        doc,
        "Video named: [Team Name]_Video_Pitch (use Zenith_Video_Pitch unless Idealize team name differs).",
    )
    add_bullet(
        doc,
        "GitHub README must cover: project purpose, tech stack (match proposal), setup instructions, "
        "core features including AI agent workflow.",
    )

    add_heading(doc, "2. Truth checks (do not oversell)", 1)
    add_bullet(
        doc,
        "Login is password-only (APP_PASSWORD). Say “login password,” not username + password.",
    )
    add_bullet(
        doc,
        "RAG / vector database is still in build — say “next milestone,” not “already live.”",
    )
    add_bullet(
        doc,
        "Cheque Print exists on preview (browser print). Full bank stationery cheque print, "
        "dedicated clearing tracker page, and owner WhatsApp clearing alerts are roadmap.",
    )
    add_bullet(
        doc,
        "This recording assumes live Gemini, OpenAI Bundling Assistant / Analyst, "
        "WhatsApp Cloud API via Cloudflare tunnel, and agents running.",
    )

    add_heading(doc, "3. Pre-demo checklist (before Record)", 1)
    add_numbered(doc, "python app.py running; Cloudflare tunnel pointing to WhatsApp webhook.")
    add_numbered(
        doc,
        "Two sample invoice photos ready in WhatsApp inbox: Invoice A = known dealer; "
        "Invoice B = new/unregistered supplier.",
    )
    add_numbered(doc, "At least one merchant bank account ready on Bank Balance.")
    add_numbered(doc, "One dealer with casual_days / impossible_days set (to explain float).")
    add_numbered(doc, "OpenAI credits + Gemini key working; avoid USE_FAKE_AI unless emergency.")
    add_numbered(doc, "Browser in English; light mode first; camera PiP visible.")

    add_heading(doc, "4. Timed schedule (~4:50)", 1)
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    headers = ["Time", "Screen", "Goal"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("0:00–0:20", "Title / login", "Problem + login"),
        ("0:20–0:45", "Bank Balance", "Add/select bank; ready for invoices"),
        ("0:45–1:50", "Invoices → WhatsApp → AI → verify (2 invoices)", "Vision agent + dealer difference"),
        ("1:50–3:25", "Cheques tab", "Options, batching, holiday float, AI suggestions"),
        ("3:25–3:55", "Bank Balance + Reports", "Cheques about to clear + analyst report"),
        ("3:55–4:25", "Dark mode, languages, Guide chat", "Local SME UX"),
        ("4:25–4:50", "Roadmap close", "RAG, print, clearing page, WhatsApp alerts"),
    ]
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = val

    add_heading(doc, "5. Cheques tab — every option to name or show", 1)
    doc.add_paragraph(
        "On the dealer Cheques page, briefly point at as many of these as time allows "
        "(name them even if you do not click every control):"
    )
    options = [
        "LKR ceiling (max per cheque)",
        "Select invoices / Select all",
        "Compute bundles",
        "One cheque per invoice",
        "Proposed cheque cards (stated date, settlement, fund-by / keep money until, Extra days gained, interbank badge)",
        "Drag invoices between cheques",
        "Move to… dropdown (other cheque / new cheque)",
        "Add cheque",
        "Right-click Split → separate cheques / same cheque / undo split (red INV · 1, INV · 2)",
        "Edit stated cheque date on a card",
        "Auto-optimize bundles (reviewer loop)",
        "Bundling Assistant chat (mic / send / stop / mute / speak / reset / hide)",
        "Apply reviewer suggestions",
        "Preview & write cheques → bank account select → cheque numbers → Print → Commit to database",
        "Committed cheques table + pending verification links (if visible)",
    ]
    for opt in options:
        add_numbered(doc, opt)

    add_heading(doc, "6. AI agents — say inputs → steps → outputs", 1)
    doc.add_paragraph(
        "Open Category requirement: clearly show the AI agent working. "
        "Call out at least Vision and Bundling Assistant."
    )
    agents = doc.add_table(rows=4, cols=4)
    agents.style = "Table Grid"
    for i, h in enumerate(["Agent", "Input", "Steps", "Output"]):
        agents.rows[0].cells[i].text = h
    agent_rows = [
        (
            "Vision (Gemini)",
            "Invoice photo (WhatsApp / upload)",
            "Extract + structure fields; human verifies",
            "Invoice #, dates, amount, supplier, line items on verify screen",
        ),
        (
            "Bundling Assistant (OpenAI + tools)",
            "Current bundles + dealer context + chat",
            "Tool calls into Python bundling / guardrails (no invented dates)",
            "Reply + updated cheque groups; Apply suggestions",
        ),
        (
            "Analyst (OpenAI)",
            "Committed cheques / metrics after commit",
            "Generate narrative report",
            "Reports page markdown analysis",
        ),
    ]
    for r_i, row in enumerate(agent_rows, start=1):
        for c_i, val in enumerate(row):
            agents.rows[r_i].cells[c_i].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        "Flow: WhatsApp photo → Vision agent → human verify/save → Bundling Assistant + Python "
        "bundling/guardrails → Preview/Commit → Analyst report."
    )

    add_heading(doc, "7. Spoken script (say this while clicking)", 1)
    doc.add_paragraph(
        "Read the italic lines aloud. Follow the bold “On screen” cues. Keep total under 5:00."
    )

    add_heading(doc, "0:00–0:20 — Intro + login", 2)
    add_quote(
        doc,
        "Hello, we are team Zenith. Our app helps Sri Lankan SMEs turn supplier invoices into "
        "smarter post-dated cheques—using Sri Lanka’s weekend and CBSL holiday clearing lag as legal float. "
        "I’m logging in with our app password on this machine. After login, the merchant is ready to "
        "configure banking and receive invoices.",
    )
    add_action(doc, "Show login page → enter password → Invoices dashboard.")

    add_heading(doc, "0:20–0:45 — Bank details first", 2)
    add_quote(
        doc,
        "First the merchant adds their bank account details on Bank Balance—nickname, bank name, branch, "
        "and opening balance—and can keep several accounts. Cheques will later be written from a chosen account. "
        "Once banking is set, they are ready to receive invoices.",
    )
    add_action(doc, "Open Bank Balance; add or select an account; show balance cards.")

    add_heading(doc, "0:45–1:50 — WhatsApp → AI → two invoices", 2)
    add_quote(
        doc,
        "Cashiers send invoice photos on WhatsApp. Photos land in the WhatsApp inbox first. "
        "I open Invoices, open WhatsApp photos, and tap Send to AI. "
        "That’s our vision agent: input is the photo, steps are Gemini extraction plus checks, "
        "output is structured fields on the verify screen. "
        "Invoice one—supplier already registered. I confirm invoice number, date, amount, credit period, and dealer. "
        "Here we set supplier rules: casual days—extra business days the dealer usually allows—and impossible days "
        "like Sunday, plus strictness. Python uses that when proposing cheque dates. "
        "Invoice two—new supplier. The app asks us to register the dealer once—name, bank, pay-from account—then save. "
        "Same dealer won’t be duplicated next time; same dealer also cannot reuse the same invoice number. "
        "Both invoices are verified and ready to bundle.",
    )
    add_action(
        doc,
        "WhatsApp inbox → Send to AI → verify Invoice A (existing dealer + casual/impossible days). "
        "Then Invoice B (new dealer form) → save. If short on time, fully verify A and quickly show B’s pending-dealer form.",
    )

    add_heading(doc, "1:50–3:25 — Cheques: options, batching, holiday loop, AI", 2)
    add_quote(
        doc,
        "Now Cheques. I pick the dealer. "
        "Here are the bundling tools: set the LKR ceiling; tick invoices; Compute bundles or One per invoice; "
        "drag or Move to another cheque; Add cheque; right-click Split for part payments shown in red as invoice ·1, ·2; "
        "edit the stated date; run Auto-optimize; and chat with the Bundling Assistant. "
        "Batching: Python packs invoices under the ceiling and proposes stated dates from due dates plus casual days, "
        "then rolls weekends and CBSL holidays to true settlement and funding dates. That’s why this is built for "
        "Sri Lanka—the holiday loop can give extra float days. On each cheque we show Extra days gained. "
        "I’ll ask the assistant to improve liquidity—or open Auto-optimize. Input is the current bundles and context; "
        "steps are tool calls into our deterministic Python bundling and guardrails—no invented dates; "
        "output is a reply plus updated groups. If the reviewer suggests changes, I click Apply suggestions. "
        "We’re also designing a RAG vector store—still in build—so past successful cheque groupings for this dealer "
        "can be retrieved next time to guide future batches.",
    )
    add_action(
        doc,
        "Open dealer Cheques → Compute → point at Extra days → Auto-optimize or chat → Apply if shown → optional split.",
    )

    add_heading(doc, "3:25–3:55 — Bank section + report", 2)
    add_quote(
        doc,
        "Back on Bank Balance, for this account we see cheques from this account and the liquidity timetable—"
        "fund-by dates and days gained—so the owner sees what’s about to clear. "
        "After we preview—choose the paying bank—and commit, the Analyst agent writes a report. "
        "On Reports we open the generated analysis. Today we print from the cheque preview; "
        "full cheque stationery print is still being finished.",
    )
    add_action(
        doc,
        "Bank Balance → cheques list + timetable. Preview → select bank → Print or Commit. Open Reports.",
    )

    add_heading(doc, "3:55–4:25 — Dark mode, Guide, languages", 2)
    add_quote(
        doc,
        "For local merchants: dark mode here; the Zenith Guide chatbot answers how-to questions; "
        "and languages—English, Sinhala, and Tamil—so the same workflow works for local business owners.",
    )
    add_action(
        doc,
        "Toggle dark mode; open Guide (one short question); switch language briefly, then back to English.",
    )

    add_heading(doc, "4:25–4:50 — Improvements + close", 2)
    add_quote(
        doc,
        "Next improvements: finish cheque print layouts; a separate clearing tracker for cheques about to go; "
        "and WhatsApp notifications to the owner when a clearing date is near. Plus the RAG memory for per-dealer history. "
        "Zenith: WhatsApp capture, AI extraction, Sri Lanka holiday-aware bundling, and multi-bank cash timing—"
        "in one local app. Thank you.",
    )
    add_action(doc, "Stay on a clean screen (Reports or Cheques). End recording.")

    add_heading(doc, "8. If something fails mid-recording", 1)
    add_quote(
        doc,
        "While live AI recovers, the same Python bundling still runs on Compute bundles—"
        "and demo mode can stand in for chat.",
    )

    add_heading(doc, "9. README reminder (submit with the repo)", 1)
    doc.add_paragraph("Before upload, confirm README.md includes:")
    add_bullet(doc, "Project purpose (Sri Lankan SME cheque / invoice liquidity).")
    add_bullet(
        doc,
        "Tech stack matching the Idealize proposal (Flask, SQLite, Gemini, OpenAI, WhatsApp, etc.).",
    )
    add_bullet(
        doc,
        "Setup: copy .env.example, APP_PASSWORD, API keys, DB init, run app, Cloudflare/WhatsApp webhook.",
    )
    add_bullet(
        doc,
        "Core features + AI agent workflow (Vision → verify → Bundling Assistant → Analyst).",
    )

    add_heading(doc, "10. Extra features to mention if you have 10–15 seconds spare", 1)
    add_bullet(doc, "Duplicate protection: same dealer name reused; same invoice number blocked per dealer.")
    add_bullet(doc, "Multi-bank: pay-from account on dealer; bank select when writing cheques.")
    add_bullet(doc, "Deposit alerts / planned deposits on Bank Balance.")
    add_bullet(doc, "Interbank +1 business day when merchant bank ≠ supplier bank.")

    doc.add_paragraph()
    end = doc.add_paragraph("End of presenter guide — practice once with a timer before the final take.")
    if end.runs:
        end.runs[0].italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
