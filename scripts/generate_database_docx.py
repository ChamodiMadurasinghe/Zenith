"""Generate Zenith database diagrams + table catalog as a Word document."""

from __future__ import annotations

import base64
import json
import re
import tempfile
import zlib
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SCHEMA_PATH = ROOT / "database" / "schema.sql"
MD_PATH = ROOT / "database" / "DATABASE.md"
OUTPUT = ROOT / "docs" / "Zenith-Database-Reference.docx"

MERMAID_INK_TIMEOUT = 45


def sanitize_mermaid(code: str) -> str:
    """mermaid.ink rejects some ER tokens used in DATABASE.md."""
    return code.replace("PK_FK", "PK")


def mermaid_pako(code: str) -> str:
    state = {
        "code": code.strip() + "\n",
        "mermaid": {"theme": "default"},
        "autoSync": True,
        "updateDiagram": True,
    }
    raw = json.dumps(state, separators=(",", ":")).encode("utf-8")
    compressed = zlib.compress(raw, 9)
    return base64.urlsafe_b64encode(compressed).decode("ascii").rstrip("=")


def mermaid_live_url(code: str) -> str:
    return f"https://mermaid.live/edit#pako:{mermaid_pako(code)}"


def mermaid_ink_url(code: str) -> str:
    return f"https://mermaid.ink/img/pako:{mermaid_pako(code)}"


def extract_mermaid_blocks(md: str) -> list[tuple[str, str]]:
    """Return (heading, mermaid_source) from DATABASE.md."""
    heading = "Diagram"
    blocks: list[tuple[str, str]] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("## "):
            heading = line[3:].strip()
        if line.strip() == "```mermaid":
            i += 1
            body: list[str] = []
            while i < len(lines) and lines[i].strip() != "```":
                body.append(lines[i])
                i += 1
            blocks.append((heading, "\n".join(body).strip()))
        i += 1
    return blocks


def parse_md_column_notes(md: str) -> dict[tuple[str, str], str]:
    notes: dict[tuple[str, str], str] = {}
    table_name: str | None = None
    in_col_table = False
    for line in md.splitlines():
        m = re.match(r"^### `([a-z_]+)`", line)
        if m:
            table_name = m.group(1)
            in_col_table = False
            continue
        if table_name and re.match(r"^\| Column \| Type \| Notes \|", line, re.I):
            in_col_table = True
            continue
        if in_col_table:
            if not line.startswith("|"):
                in_col_table = False
                continue
            if re.match(r"^\|\s*-+", line):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cells) >= 3:
                col = cells[0].strip().strip("`")
                notes[(table_name, col)] = cells[2].strip()
    return notes


def parse_md_table_intros(md: str) -> dict[str, str]:
    intros: dict[str, str] = {}
    table_name: str | None = None
    buf: list[str] = []
    for line in md.splitlines():
        m = re.match(r"^### `([a-z_]+)`", line)
        if m:
            if table_name and buf:
                intros[table_name] = " ".join(buf).strip()
            table_name = m.group(1)
            buf = []
            continue
        if table_name:
            if line.startswith("|") or line.startswith("### ") or line.startswith("## "):
                if table_name and buf:
                    intros[table_name] = " ".join(buf).strip()
                    buf = []
                if line.startswith("## ") or line.startswith("### "):
                    table_name = m.group(1) if line.startswith("### `") else None
                continue
            if line.strip() and not line.startswith("---"):
                buf.append(line.strip())
    if table_name and buf:
        intros[table_name] = " ".join(buf).strip()
    return intros


def _split_sql_columns(body: str) -> list[str]:
    parts: list[str] = []
    current: list[str] = []
    depth = 0
    for ch in body:
        if ch == "(":
            depth += 1
            current.append(ch)
        elif ch == ")":
            depth -= 1
            current.append(ch)
        elif ch == "," and depth == 0:
            parts.append("".join(current).strip())
            current = []
        else:
            current.append(ch)
    if current and "".join(current).strip():
        parts.append("".join(current).strip())
    return parts


def parse_schema(sql: str) -> tuple[list[dict], list[dict]]:
    tables: list[dict] = []
    indexes: list[dict] = []

    for m in re.finditer(
        r"CREATE TABLE (\w+)\s*\((.*?)\)\s*;",
        sql,
        re.DOTALL | re.IGNORECASE,
    ):
        name = m.group(1)
        body = m.group(2)
        columns: list[dict] = []
        table_fks: dict[str, str] = {}
        for raw in _split_sql_columns(body):
            upper = raw.upper()
            fk_m = re.search(
                r"FOREIGN KEY\s*\((\w+)\)\s*REFERENCES\s+(\w+)\s*\((\w+)\)",
                raw,
                re.I,
            )
            if fk_m:
                table_fks[fk_m.group(1)] = f"{fk_m.group(2)}({fk_m.group(3)})"
                continue
            if upper.startswith("PRIMARY KEY") or upper.startswith("UNIQUE"):
                continue
            col_m = re.match(r"(\w+)\s+(INTEGER|TEXT|REAL|BLOB|NUMERIC)\b(.*)$", raw, re.I)
            if not col_m:
                continue
            col, typ, rest = col_m.group(1), col_m.group(2).upper(), col_m.group(3)
            rest_u = rest.upper()
            keys: list[str] = []
            if "PRIMARY KEY" in rest_u:
                keys.append("PK")
            if re.search(r"\bUNIQUE\b", rest_u):
                keys.append("UNIQUE")
            inline_fk = re.search(r"REFERENCES\s+(\w+)\s*\((\w+)\)", rest, re.I)
            if inline_fk:
                keys.append(f"FK → {inline_fk.group(1)}({inline_fk.group(2)})")
            nullable = "NOT NULL" not in rest_u
            if "PRIMARY KEY" in rest_u:
                nullable = False
            default_m = re.search(r"DEFAULT\s+(\([^)]+\)|'[^']*'|\S+)", rest, re.I)
            default = default_m.group(1) if default_m else ""
            columns.append(
                {
                    "name": col,
                    "type": typ,
                    "keys": keys,
                    "nullable": nullable,
                    "default": default,
                    "rest": rest.strip(),
                }
            )
        for col in columns:
            if col["name"] in table_fks:
                ref = f"FK → {table_fks[col['name']]}"
                if ref not in col["keys"]:
                    col["keys"].append(ref)
        tables.append({"name": name, "columns": columns})

    for m in re.finditer(
        r"CREATE (UNIQUE )?INDEX (\w+)\s+ON (\w+)\s*\((.*?)\)\s*;",
        sql,
        re.DOTALL | re.IGNORECASE,
    ):
        unique = bool(m.group(1))
        indexes.append(
            {
                "name": m.group(2),
                "unique": unique,
                "table": m.group(3),
                "columns": re.sub(r"\s+", " ", m.group(4)).strip(),
            }
        )
        if unique:
            for tbl in tables:
                if tbl["name"] != m.group(3):
                    continue
                cols = [c.strip() for c in m.group(4).split(",")]
                if len(cols) == 1:
                    for col in tbl["columns"]:
                        if col["name"] == cols[0] and "UNIQUE" not in col["keys"]:
                            col["keys"].append("UNIQUE")

    return tables, indexes


def fetch_mermaid_image(code: str, dest_stem: Path) -> Path | None:
    url = mermaid_ink_url(code)
    try:
        resp = requests.get(url, timeout=MERMAID_INK_TIMEOUT)
    except requests.RequestException:
        return None
    if resp.status_code != 200 or len(resp.content) < 200:
        return None
    data = resp.content
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        path = dest_stem.with_suffix(".png")
    elif data[:3] == b"\xff\xd8\xff":
        path = dest_stem.with_suffix(".jpg")
    else:
        ctype = (resp.headers.get("content-type") or "").lower()
        if "jpeg" in ctype or "jpg" in ctype:
            path = dest_stem.with_suffix(".jpg")
        elif "png" in ctype:
            path = dest_stem.with_suffix(".png")
        else:
            return None
    path.write_bytes(data)
    return path


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def fill_header_row(row, headers: list[str]) -> None:
    for i, h in enumerate(headers):
        row.cells[i].text = h
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(row.cells[i], "E8EDF4")


def add_grid(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    fill_header_row(table.rows[0], headers)
    for r_i, values in enumerate(rows):
        for c_i, val in enumerate(values):
            table.rows[r_i + 1].cells[c_i].text = val
    doc.add_paragraph("")


def column_null_default(col: dict) -> str:
    bits: list[str] = []
    bits.append("NULL ok" if col["nullable"] else "NOT NULL")
    if col["default"]:
        bits.append(f"default {col['default']}")
    return "; ".join(bits)


def build_document() -> tuple[Document, list[tuple[str, str, bool]]]:
    sql = SCHEMA_PATH.read_text(encoding="utf-8")
    md = MD_PATH.read_text(encoding="utf-8")
    tables, indexes = parse_schema(sql)
    notes = parse_md_column_notes(md)
    intros = parse_md_table_intros(md)
    diagrams = extract_mermaid_blocks(md)

    doc = Document()
    title = doc.add_heading("Zenith Database Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("SQLite schema, diagrams, and field catalog")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True

    doc.add_paragraph(
        "This document describes the Zenith / ChequeMate SQLite database "
        "(database/invoice_cheque.db). Schema source: database/schema.sql. "
        "Rebuild with python scripts/init_db.py."
    )
    doc.add_paragraph(
        "The entity-relationship diagram is large and may not fit a printed page. "
        "Use the mermaid.live link under that figure to zoom, pan, and export PNG/SVG for teammates."
    )

    doc.add_heading("Shareable diagram links", 1)
    diagram_results: list[tuple[str, str, bool]] = []
    tmpdir = Path(tempfile.mkdtemp(prefix="zenith-mermaid-"))

    for heading, code in diagrams:
        code = sanitize_mermaid(code)
        live = mermaid_live_url(code)
        dest_stem = tmpdir / re.sub(r"[^a-z0-9]+", "-", heading.lower()).strip("-")
        img_path = fetch_mermaid_image(code, dest_stem)
        embedded = img_path is not None
        diagram_results.append((heading, live, embedded))

        doc.add_heading(heading, 2)
        if img_path:
            doc.add_picture(str(img_path), width=Inches(6.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(heading)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            doc.add_paragraph(
                "Diagram image could not be downloaded (mermaid.ink). "
                "Open the live link below to view and export it."
            )
            src = doc.add_paragraph()
            run = src.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(8)

        p = doc.add_paragraph("Open, zoom, and share: ")
        add_hyperlink(p, "mermaid.live — " + heading, live)

    doc.add_heading("How invoices link to cheques", 1)
    doc.add_paragraph(
        "A. Simple link — invoices.cheque_id. One invoice points at one primary cheque (classic bundling)."
    )
    doc.add_paragraph(
        "B. Split link — cheque_invoice_allocation. One invoice can fund multiple cheques "
        "(or partial amounts). Committed cheques use both: cheque_id on the invoice plus allocation rows for splits."
    )

    doc.add_heading("Relationship summary", 1)
    add_grid(
        doc,
        ["From", "To", "Cardinality", "Meaning"],
        [
            ["user", "user_bank_account", "1:N", "Merchant owns bank accounts"],
            ["user", "invoices", "1:N", "Merchant owns invoices"],
            ["dealers", "invoices", "1:N", "Supplier has many invoices"],
            ["invoices", "item", "1:N", "Invoice has line items"],
            ["invoices", "cheque", "N:1 (nullable)", "Many invoices on one cheque"],
            ["invoices", "cheque via allocation", "N:M", "Invoice can split across cheques"],
            ["cheque", "user_bank_account", "N:1", "Cheque drawn from one account"],
            ["cheque", "deposit_timetable", "1:1-ish", "Liquidity tracking per cheque"],
            ["dealers", "bundle_drafts", "1:1", "One active bundling session per supplier"],
            ["whatsapp_inbox", "invoices", "N:1", "Image becomes an invoice"],
        ],
    )

    doc.add_heading("Table catalog", 1)
    doc.add_paragraph(
        "Each table lists Column, SQLite type, Key (PK / FK / UNIQUE), Null/default, and notes. "
        "Types come from schema.sql. Notes come from database/DATABASE.md."
    )

    for tbl in tables:
        doc.add_heading(tbl["name"], 2)
        intro = intros.get(tbl["name"])
        if intro:
            doc.add_paragraph(intro)
        rows = []
        for col in tbl["columns"]:
            key = ", ".join(col["keys"]) if col["keys"] else ""
            note = notes.get((tbl["name"], col["name"]), "")
            rows.append(
                [
                    col["name"],
                    col["type"],
                    key,
                    column_null_default(col),
                    note,
                ]
            )
        add_grid(doc, ["Column", "Type", "Key", "Null / default", "Notes"], rows)

    doc.add_heading("app_settings keys", 2)
    doc.add_paragraph("Key-value rows stored in app_settings (not columns).")
    add_grid(
        doc,
        ["setting_key", "Default", "Purpose"],
        [
            ["min_cash_buffer_lkr", "500000", "Minimum safe balance"],
            ["default_bank_acc_id", "1", "Default account for Cash Flow / bundling"],
        ],
    )

    doc.add_heading("Indexes", 1)
    idx_rows = []
    for idx in indexes:
        kind = "UNIQUE" if idx["unique"] else "INDEX"
        idx_rows.append([idx["name"], idx["table"], idx["columns"], kind])
    add_grid(doc, ["Index", "Table", "Columns", "Kind"], idx_rows)

    doc.add_heading("Liquidity formula (conceptual)", 1)
    formula = doc.add_paragraph()
    run = formula.add_run(
        "usable_funds = available_balance + overdraft_limit + planned_deposits\n"
        "outflows     = deposit_timetable rows (status = pending)\n"
        "safe_to_issue = usable_funds - outflows - min_cash_buffer_lkr"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    doc.add_paragraph(
        "Date calculations for cheque clearance and funding use cbsl_bank_holidays "
        "plus each dealer’s casual_days and impossible_days."
    )

    return doc, diagram_results


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc, results = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print()
    print("Shareable mermaid.live links (open these if a diagram is too large for print/PDF):")
    for heading, url, embedded in results:
        flag = "embedded" if embedded else "link only"
        print(f"  [{flag}] {heading}")
        print(f"    {url}")
        print()


if __name__ == "__main__":
    main()
