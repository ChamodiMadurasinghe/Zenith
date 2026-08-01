"""Generate Zenith teammate implementation guide as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Zenith-Implementation-Plan.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def build_document():
    doc = Document()

    # Title
    title = doc.add_heading("Zenith Cheque Writing Web App", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Teammate Implementation Guide")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph(
        "This document explains how to set up, understand, and build the Zenith cheque-writing "
        "application. It is a local, single-user Python (Flask) web app for Sri Lankan businesses, "
        "with three Gemini AI agents and deterministic Python logic for bundling, guardrails, and cash-flow planning."
    )

    # 1. Overview
    add_heading(doc, "1. Project Overview", 1)
    doc.add_paragraph(
        "Zenith helps a merchant manage supplier invoices, verify scanned invoice images, bundle "
        "invoices into cheques under an LKR ceiling, print cheques, and know when to deposit money "
        "into the bank to cover upcoming cheque clearances."
    )
    add_bullet(doc, "Runs locally on one machine (Flask + SQLite).", "Deployment: ")
    add_bullet(doc, "One merchant account with a password login screen.", "Users: ")
    add_bullet(doc, "Python (Flask) for all backend logic; HTML/JS templates for the UI.", "Backend: ")
    add_bullet(doc, "Google Gemini for vision ingestion, chat assistant, and analytics narrative.", "AI: ")

    # 2. Architecture
    add_heading(doc, "2. High-Level Architecture", 1)
    doc.add_paragraph(
        "The app has four main areas: (1) Ingestion & verification, (2) Bundling & cheque printing, "
        "(3) Analyst dashboard, and (4) Cash Flow deposit timing. Agents never write to the database "
        "directly — Python is the sole orchestrator and database writer."
    )

    add_heading(doc, "Agent Responsibilities", 2)
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["Agent", "Role", "Gemini Usage", "Python Owns"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        (
            "Agent 1 — Ingestion",
            "Extract data from invoice images",
            "Image → structured JSON (invoice #, supplier, amount, line items)",
            "File storage, dealer lookup, routing to review screen",
        ),
        (
            "Agent 2 — Assistant",
            "Human-in-the-loop help",
            "Dealer setup suggestions; bundling chat explaining dates/groups",
            "All DB writes, guardrail enforcement, applying confirmed changes",
        ),
        (
            "Agent 3 — Analyst",
            "Business insights",
            "Markdown narrative from pre-computed metrics",
            "SQL aggregation, async job scheduling",
        ),
    ]
    for r, row_data in enumerate(rows, start=1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        "Critical rule: Agent 2 proposes changes as structured actions (e.g. set_cheque_date). "
        "Python validates every action through core/guardrails.py before updating state or preview."
    )

    # 3. Workflow phases
    add_heading(doc, "3. Workflow Phases", 1)

    add_heading(doc, "Phase 1 — Ingestion & Human-in-the-Loop Verification", 2)
    add_numbered(doc, "Cashier uploads an invoice image (JPG/PNG).")
    add_numbered(doc, "Agent 1 (Gemini vision) extracts invoice number, supplier name, amount, and line items.")
    add_numbered(doc, "Python checks if the dealer exists in the database.")
    add_numbered(doc, "If dealer is unknown: Agent 2 pre-fills a dealer setup form; user confirms.")
    add_numbered(doc, "Review screen shows extracted fields; user edits and clicks Verify & Save.")
    add_numbered(doc, "Python saves invoice + item rows and sets is_invoice_verified = 1.")

    add_heading(doc, "Phase 2 — Bundling, Negotiation & Guardrails", 2)
    add_numbered(doc, "Merchant selects a dealer, verified invoices, and an LKR ceiling.")
    add_numbered(doc, "Python bundling engine groups invoices into cheque batches (greedy bin-pack by due date).")
    add_numbered(doc, "Agent 2 chat explains the proposed dates and groups conversationally.")
    add_numbered(doc, "User can adjust dates/groups in chat; Agent 2 emits proposed_actions JSON.")
    add_numbered(doc, "Python guardrails check: ceiling not broken, no bank holidays, no dealer impossible_days, balance OK.")
    add_numbered(doc, "On PASS: show cheque preview; user prints; Python commits cheque rows and links invoices.")
    add_numbered(doc, "On FAIL: Agent 2 guides user to adjust; chat loop continues.")

    add_heading(doc, "Phase 3 — Analyst Engine (Async)", 2)
    add_numbered(doc, "After cheque commit, a background job runs.")
    add_numbered(doc, "Python queries bank deposits, outstanding liabilities, and clearance trends.")
    add_numbered(doc, "Agent 3 generates markdown inventory/cash-flow commentary.")
    add_numbered(doc, "Report is cached and shown on the Analytics page.")

    add_heading(doc, "Cash Flow — When to Deposit Money", 2)
    doc.add_paragraph(
        "A separate /cash-flow page (Python-only, no Gemini) shows deposit alerts and a 60-day "
        "balance timeline. It uses available_balance, upcoming cheque clearances, planned deposits, "
        "and min_cash_buffer_lkr from app_settings."
    )
    add_bullet(doc, "Deposit alerts — e.g. \"Deposit LKR 450,000 by 2025-06-18 to cover cheques clearing on 2025-06-20\"")
    add_bullet(doc, "Day-by-day timeline — projected balance; red when below minimum buffer")
    add_bullet(doc, "Quick actions — update balance, add planned deposit, mark deposit as done")

    # 4. Project structure
    add_heading(doc, "4. Project Structure", 1)
    structure = """Zenith-cursor/
├── .env.example              # Copy to .env and add API keys + password
├── app.py                    # Flask entry point (wire all blueprints)
├── config.py                 # Loads .env via python-dotenv
├── requirements.txt
├── agents/
│   ├── base.py               # Gemini client wrapper
│   ├── ingestion.py          # Agent 1
│   ├── assistant.py          # Agent 2
│   └── analyst.py            # Agent 3
├── core/
│   ├── auth.py               # @login_required, password verify
│   ├── bundling.py           # LKR ceiling grouping + date optimization
│   ├── guardrails.py         # Holiday/ceiling/balance checks
│   ├── cash_flow.py          # Deposit timing projections
│   ├── amounts.py            # LKR amount-in-words for cheques
│   └── dates.py              # Business day helpers
├── db/
│   ├── connection.py
│   └── repositories.py
├── routes/
│   ├── auth.py               # /login, /logout
│   ├── ingestion.py          # Phase 1
│   ├── bundling.py           # Phase 2
│   ├── analytics.py          # Phase 3
│   └── cash_flow.py          # Deposit timing
├── templates/                # HTML/Jinja pages
├── static/css, static/js
├── database/
│   ├── schema.sql
│   ├── seed.sql
│   ├── DATABASE.md           # Full table reference
│   └── invoice_cheque.db
├── scripts/
│   └── init_db.py            # Rebuild database
└── storage/invoices/         # Uploaded images (gitignored)"""
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 5. Setup
    add_heading(doc, "5. Getting Started (For Every Teammate)", 1)

    add_heading(doc, "Prerequisites", 2)
    add_bullet(doc, "Python 3.10+")
    add_bullet(doc, "A Google Gemini API key (GEMINI_API_KEY)")

    add_heading(doc, "Setup Steps", 2)
    add_numbered(doc, "Clone the repository and open a terminal in the project root.")
    add_numbered(doc, "Create a virtual environment: python -m venv .venv")
    add_numbered(doc, "Activate it: .venv\\Scripts\\activate (Windows) or source .venv/bin/activate (Mac/Linux)")
    add_numbered(doc, "Install dependencies: pip install -r requirements.txt")
    add_numbered(doc, "Copy .env.example to .env and fill in:")
    env_items = [
        "FLASK_SECRET_KEY — random secret string",
        "GEMINI_API_KEY — your Gemini API key",
        "APP_PASSWORD — the password you will use to log in",
    ]
    for item in env_items:
        add_bullet(doc, item)
    add_numbered(doc, "Initialize the database: python scripts/init_db.py")
    add_numbered(doc, "Run the app: python app.py (or flask run)")
    add_numbered(doc, "Open http://127.0.0.1:5000/login and enter your APP_PASSWORD.")

    doc.add_paragraph(
        "To reset the database to a clean state: run python scripts/init_db.py again. "
        "This wipes all data and re-applies schema + seed."
    )

    # 6. Database
    add_heading(doc, "6. Database Structure", 1)
    doc.add_paragraph(
        "SQLite database at database/invoice_cheque.db. Full column-by-column reference is in "
        "database/DATABASE.md. Summary of tables:"
    )
    db_table = doc.add_table(rows=13, cols=3)
    db_table.style = "Table Grid"
    db_headers = ["Table", "Status", "Purpose"]
    for i, h in enumerate(db_headers):
        db_table.rows[0].cells[i].text = h
    db_rows = [
        ("user", "Kept", "Single merchant; password_hash for login"),
        ("user_bank_account", "Kept", "Your bank accounts + available_balance"),
        ("dealers", "Kept", "Supplier registry"),
        ("dealers_bank_account", "Kept", "Supplier bank details"),
        ("cheque", "Kept", "Issued cheques"),
        ("invoices", "Kept", "Invoice headers; is_invoice_verified, cheque_id"),
        ("item", "Kept", "Invoice line items"),
        ("cbsl_bank_holidays", "Kept", "Sri Lanka bank holidays"),
        ("bank_deposits", "New", "Actual money deposited"),
        ("planned_deposits", "New", "Future deposits you plan to make"),
        ("app_settings", "New", "min_cash_buffer_lkr, default_bank_acc_id"),
        ("analyst_reports", "New", "Cached Agent 3 markdown reports"),
    ]
    for r, row_data in enumerate(db_rows, start=1):
        for c, val in enumerate(row_data):
            db_table.rows[r].cells[c].text = val

    # 7. Auth
    add_heading(doc, "7. Authentication", 1)
    add_bullet(doc, "/login — public; password form only (no username field)")
    add_bullet(doc, "All other routes — protected with @login_required")
    add_bullet(doc, "/logout — clears session")
    add_bullet(doc, "Password stored as werkzeug hash in user.password_hash")
    add_bullet(doc, "Initial password comes from APP_PASSWORD in .env; init_db.py hashes it on DB creation")

    # 8. Key routes
    add_heading(doc, "8. Key Routes", 1)
    routes_table = doc.add_table(rows=16, cols=3)
    routes_table.style = "Table Grid"
    for i, h in enumerate(["Route", "Method", "Description"]):
        routes_table.rows[0].cells[i].text = h
    route_rows = [
        ("/login", "GET/POST", "Password gate"),
        ("/logout", "POST", "End session"),
        ("/", "GET", "Dashboard — upload + recent invoices"),
        ("/upload", "POST", "Save image → Agent 1 → dealer lookup"),
        ("/review/<draft_id>", "GET", "Review extracted invoice"),
        ("/review/<draft_id>/verify", "POST", "Save verified invoice"),
        ("/dealer/setup", "GET/POST", "New dealer registration"),
        ("/bundling", "GET", "Dealer picker + invoice selection"),
        ("/api/chat/*", "POST", "Agent 2 bundling chat"),
        ("/cheque/preview", "GET", "Cheque print preview"),
        ("/cash-flow", "GET", "Deposit timing timeline + alerts"),
        ("/cash-flow/balance", "POST", "Update account balance"),
        ("/cash-flow/planned", "POST", "Add planned deposit"),
        ("/analytics", "GET", "Agent 3 markdown report"),
    ]
    for r, row_data in enumerate(route_rows, start=1):
        for c, val in enumerate(row_data):
            routes_table.rows[r].cells[c].text = val

    # 9. Implementation order
    add_heading(doc, "9. Suggested Implementation Order", 1)
    doc.add_paragraph("Build in vertical slices so each phase is testable before moving on:")
    order = [
        "Scaffold + DB rebuild — Flask app, init_db.py, schema/seed, DATABASE.md, .env.example",
        "Auth — login.html, core/auth.py, /login + /logout, protect all routes",
        "Cash Flow page — core/cash_flow.py + /cash-flow UI (works with seed data, no Gemini)",
        "Phase 1 — Upload + Agent 1 + review screen + verify save",
        "Phase 2 core — Bundling engine + guardrails (unit-testable without Gemini)",
        "Phase 2 UI — Dealer picker, chat, cheque preview, print commit",
        "Phase 3 — Async Agent 3 analyst, analytics page",
        "app.py — Wire all blueprints and run entry point",
    ]
    for i, step in enumerate(order, start=1):
        add_numbered(doc, step)

    # 10. Task split
    add_heading(doc, "10. How to Split Work Among Teammates", 1)
    doc.add_paragraph("Suggested ownership areas (adjust based on your team size):")

    split_table = doc.add_table(rows=5, cols=3)
    split_table.style = "Table Grid"
    for i, h in enumerate(["Area", "Owner Focus", "Key Files"]):
        split_table.rows[0].cells[i].text = h
    split_rows = [
        (
            "Infrastructure & Auth",
            "Flask scaffold, config, DB layer, login, init_db",
            "app.py, config.py, db/, core/auth.py, routes/auth.py, scripts/init_db.py",
        ),
        (
            "Phase 1 — Ingestion",
            "Agent 1, upload flow, review/dealer templates",
            "agents/ingestion.py, routes/ingestion.py, templates/upload.html, review_invoice.html",
        ),
        (
            "Phase 2 — Bundling",
            "Bundling algorithm, guardrails, chat, cheque preview",
            "core/bundling.py, core/guardrails.py, agents/assistant.py, routes/bundling.py",
        ),
        (
            "Cash Flow & Analytics",
            "Deposit projections, planned deposits UI, Agent 3",
            "core/cash_flow.py, routes/cash_flow.py, agents/analyst.py, routes/analytics.py",
        ),
    ]
    for r, row_data in enumerate(split_rows, start=1):
        for c, val in enumerate(row_data):
            split_table.rows[r].cells[c].text = val

    # 11. Bundling logic
    add_heading(doc, "11. Bundling Algorithm (Python — Not LLM)", 1)
    doc.add_paragraph("core/bundling.py implements deterministic cheque batching:")
    add_numbered(doc, "Input: dealer_id, verified invoice IDs (cheque_id IS NULL), ceiling_lkr, dealer casual_days/impossible_days.")
    add_numbered(doc, "Sort invoices by due date (invoiced_date + credit_period_days).")
    add_numbered(doc, "Greedy bin-pack: group invoices where sum(total_amount) <= ceiling_lkr.")
    add_numbered(doc, "For each batch, compute cheque_date = earliest valid business day after last due date.")
    add_numbered(doc, "Skip CBSL holidays and dealer impossible_days.")
    add_numbered(doc, "Return ChequeBundle[] with invoice_ids, cheque_date, total_lkr, predicted_clearance_date.")

    # 12. Guardrails
    add_heading(doc, "12. Guardrail Checks", 1)
    add_bullet(doc, "Sum of batch amounts <= ceiling_lkr")
    add_bullet(doc, "cheque_date not on cbsl_bank_holidays")
    add_bullet(doc, "cheque_date not on dealer.impossible_days")
    add_bullet(doc, "cheque_date >= today")
    add_bullet(doc, "available_balance >= batch total (optional balance guard)")

    # 13. Dependencies
    add_heading(doc, "13. Python Dependencies", 1)
    deps = [
        "flask>=3.0 — web framework",
        "python-dotenv>=1.0 — .env loading",
        "google-generativeai>=0.8 — Gemini API",
        "markdown>=3.5 — render Agent 3 reports",
        "Pillow>=10.0 — image validation before Gemini",
    ]
    for d in deps:
        add_bullet(doc, d)

    # 14. Decisions
    add_heading(doc, "14. Key Design Decisions", 1)
    add_bullet(doc, "Local single-user with password — not multi-tenant")
    add_bullet(doc, "Python-first backend — HTML/JS is presentation only")
    add_bullet(doc, "Gemini for all three agents — single GEMINI_API_KEY")
    add_bullet(doc, "Agents never write to DB — Python orchestrates everything")
    add_bullet(doc, "Cash Flow is deterministic Python — fast, no API cost, works offline for projections")
    add_bullet(doc, "Bundling chat state in Flask session — no bundling_sessions table")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Document generated for the Zenith team. See database/DATABASE.md and .env.example for reference.")
    footer.runs[0].italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
