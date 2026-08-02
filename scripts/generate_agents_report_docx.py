"""Generate teammate-facing Word report: agents, tool calling, guardrails, user overrides."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Zenith-Agents-Workflow-Guide.docx"


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    doc.add_paragraph(text)


def bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def callout(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = val
    doc.add_paragraph()


def build():
    doc = Document()

    title = doc.add_heading("Zenith — How AI Agents Work", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Teammate Guide: Workflow, Tool Calling, Guardrails & User Control")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True

    p(
        doc,
        "This document explains, in simple terms, how Zenith uses AI agents, how those agents "
        "call tools, how Python blocks bad decisions, and how a human user can override every "
        "AI suggestion before anything is saved.",
    )

    callout(
        doc,
        "Golden rule: AI proposes. Python checks the maths. Humans confirm. Only then does the database change.",
    )

    # ------------------------------------------------------------------
    h(doc, "1. The Big Picture (30 seconds)", 1)
    p(
        doc,
        "Zenith is a local Flask app for managing supplier invoices and writing cheques. "
        "AI helps extract invoice data from photos, suggest cheque groups, and write reports. "
        "It does not get the final say.",
    )
    numbered(doc, "AI looks at data or an image and suggests something.")
    numbered(doc, "Python tools run the real calculations (amounts, dates, ceilings, holidays).")
    numbered(doc, "Guardrails reject or warn about unsafe suggestions.")
    numbered(doc, "You review, edit, or ignore the AI — then click Verify / Commit.")
    numbered(doc, "Only your confirmed action writes to SQLite.")

    # ------------------------------------------------------------------
    h(doc, "2. Agents at a Glance", 1)
    p(doc, "Think of each agent as a specialist helper with a narrow job:")

    table(
        doc,
        ["Agent", "Simple job", "Uses AI?", "Writes to DB?"],
        [
            (
                "Ingestion (Vision)",
                "Read invoice photo → fill form fields",
                "Yes (Gemini vision)",
                "No — only after you Verify",
            ),
            (
                "Anomaly checker",
                "Flag odd invoice values (rules)",
                "Mostly rules; optional AI second look",
                "No",
            ),
            (
                "Bundling Assistant",
                "Chat to regroup cheques / change dates",
                "Yes (OpenAI + tools)",
                "No — draft in session only",
            ),
            (
                "Reviewer",
                "Second opinion on liquidity of bundles",
                "Yes (OpenAI)",
                "No — Apply is optional",
            ),
            (
                "Guide",
                "Help navigate the app",
                "Yes (OpenAI)",
                "No — navigate / logout only",
            ),
            (
                "Analyst",
                "Write a cash / inventory style report",
                "Yes (OpenAI)",
                "Yes — after you commit cheques (report cache)",
            ),
            (
                "Agentic Agents 1–4",
                "Optional WhatsApp/API pipeline",
                "Vision + rules + liquidity + liaison",
                "Does not replace main verify/commit gates",
            ),
        ],
    )

    p(
        doc,
        "Naming tip: In the web UI, “Cheque Assistant” is the Bundling Assistant "
        "(tool-calling). That is different from agentic “Agent 2” (Anomaly) in the optional pipeline.",
    )

    # ------------------------------------------------------------------
    h(doc, "3. Workflow in Simple Terms", 1)

    h(doc, "3.1 Invoice capture (Ingestion)", 2)
    numbered(doc, "Cashier uploads a photo (or sends one via WhatsApp inbox).")
    numbered(doc, "Vision agent extracts invoice number, supplier, amounts, line items.")
    numbered(doc, "Anomaly rules may flag suspicious totals or missing data.")
    numbered(doc, "Review screen shows everything editable.")
    numbered(doc, "You tick “I confirm this matches the invoice” and Verify & Save.")
    numbered(doc, "Python then inserts the invoice into the database.")

    h(doc, "3.2 Cheque bundling", 2)
    numbered(doc, "You pick a dealer, verified invoices, and an LKR ceiling.")
    numbered(doc, "Python packs invoices into cheque groups (not the LLM).")
    numbered(doc, "You can chat with the Bundling Assistant to move invoices, split, or change dates.")
    numbered(doc, "Each chat change goes through tools + guardrails.")
    numbered(doc, "Optional Reviewer can suggest liquidity tweaks — you choose Apply or ignore.")
    numbered(doc, "Preview cheques. If there are warnings, you must acknowledge them to continue.")
    numbered(doc, "You assign bank account + cheque numbers and Commit — that is the DB write.")

    h(doc, "3.3 After commit", 2)
    numbered(doc, "Analyst agent writes a markdown report from Python-prepared metrics.")
    numbered(doc, "Cash Flow page (pure Python) still shows when to deposit money — independent of chat AI.")

    # ------------------------------------------------------------------
    h(doc, "4. How Tool Calling Works", 1)
    p(
        doc,
        "“Tool calling” means the AI does not invent final numbers in secret. "
        "It asks Python functions (tools) to do the work, then reads the tool results.",
    )

    h(doc, "4.1 Bundling Assistant — real function calling (main example)", 2)
    p(
        doc,
        "This is the strongest tool-calling path. The assistant uses LangChain + OpenAI "
        "function calling. It can call tools like a calculator API — but the calculator is our code.",
    )
    bullet(doc, "Prefer dry_run=True — preview first, do not lock the draft until you confirm.", "dry_run: ")
    bullet(
        doc,
        "Call apply_bundle_changes(confirm=true) or re-run a tool with dry_run=False only when the user agrees.",
        "Confirm step: ",
    )
    bullet(
        doc,
        "If a tool returns ok=false or a day-limit warning, explain it and suggest another tool call.",
        "On failure: ",
    )

    p(doc, "Tools the Bundling Assistant can call:")
    table(
        doc,
        ["Tool name", "What it does"],
        [
            ("compute_cheque_bundles", "Pack invoices under the LKR ceiling (Python greedy pack)"),
            ("divide_into_cheques", "Split into N balanced cheque groups"),
            ("move_invoice", "Move one invoice to another cheque group"),
            ("rebatch_invoice", "Move invoice and recalculate dates"),
            ("set_cheque_date", "Set a stated cheque date; Python recalculates settlement/funding"),
            ("postpone_cheque", "Push a cheque date forward by N days"),
            ("split_invoice", "Put one invoice on its own cheque"),
            ("recalculate_dates", "Re-run liquidity / date enrichment on all groups"),
            ("create_bundles", "Replace layout with explicit groups you describe"),
            ("assign_invoices", "Map invoice IDs → group numbers (+ optional dates)"),
            ("check_day_limit_risk", "Read-only check against casual daily deposit limits"),
            ("apply_bundle_changes", "Commit the last dry-run preview into the UI draft (needs confirm=true)"),
        ],
    )

    callout(
        doc,
        "Important: These tools call core/bundling.py and core/guardrails.py. The model is not allowed to invent ceilings or holiday-safe dates on its own.",
    )

    h(doc, "4.2 Legacy / other agents — structured actions (lighter tool pattern)", 2)
    p(
        doc,
        "Older chat mode and some agents do not use LangChain tools. Instead they emit JSON "
        "blocks like proposed_actions. Python still parses and runs those actions through the same guardrails.",
    )
    bullet(doc, "Reviewer — may suggest changes as structured actions; Apply is a separate user click.")
    bullet(doc, "Guide — emits navigate / logout actions only; never bundles cheques.")
    bullet(doc, "Ingestion — vision returns structured JSON fields; Python validates and shows a form.")
    bullet(doc, "Analyst — receives pre-computed metrics from Python; writes narrative markdown.")

    h(doc, "4.3 Optional agentic pipeline (Agents 1–4)", 2)
    p(
        doc,
        "When USE_AGENTIC_ORCHESTRATOR is on (e.g. WhatsApp/API path), a separate orchestrator "
        "runs Plan → Execute → Review loops over protocol tools:",
    )
    bullet(doc, "vision.extract — wrap Gemini invoice extraction")
    bullet(doc, "anomaly.audit — rule audit (+ optional AI review)")
    bullet(doc, "liquidity.forecast — Python liquidity engine (+ optional date pick)")
    bullet(doc, "liaison.draft_message / handle_reply — dealer WhatsApp style messages")
    p(
        doc,
        "Even here, final business commits in the main app still go through human verification screens.",
    )

    # ------------------------------------------------------------------
    h(doc, "5. Python Guardrails — How Bad Decisions Are Limited", 1)
    p(
        doc,
        "Guardrails live mainly in core/guardrails.py. They run after every bundle compute, "
        "chat tool mutation, review apply, manual edit, and before preview/commit.",
    )

    h(doc, "5.1 What gets checked", 2)
    bullet(doc, "No empty cheque groups")
    bullet(doc, "Cheque total cannot exceed your LKR ceiling (unless you explicitly allow exceed)")
    bullet(doc, "Cheque stated date must exist and cannot be in the past")
    bullet(doc, "Funding / clearance date cannot already be past")
    bullet(doc, "Same invoice cannot appear on two cheques")
    bullet(doc, "Invoice already linked to a committed cheque cannot be rebundled")
    bullet(doc, "Casual daily limit risk (day exposure) is audited and can warn LIMIT_BREACH_WARNING")

    h(doc, "5.2 Why this matters", 2)
    p(
        doc,
        "If the AI asks to postpone a cheque to an unsafe day, or packs more money than your ceiling, "
        "the tool path still runs Python validation. The UI shows issues. Commit can be blocked or "
        "require you to acknowledge warnings. The AI cannot silently bypass these checks.",
    )

    h(doc, "5.3 Separation of duties", 2)
    table(
        doc,
        ["Layer", "Allowed to do", "Not allowed to do"],
        [
            ("LLM / agent", "Suggest text, call tools, propose actions", "Write cheques/invoices directly to DB"),
            ("Python tools", "Pack, date maths, mutate draft state", "Skip guardrail checks"),
            ("Guardrails", "Validate / reject / warn", "Guess business intent"),
            ("Human user", "Edit, override, verify, commit", "—"),
        ],
    )

    # ------------------------------------------------------------------
    h(doc, "6. How You Override Everything the AI Decides", 1)
    p(
        doc,
        "Zenith is built as human-in-the-loop. At every major step you can change or ignore AI output.",
    )

    h(doc, "6.1 Invoice stage overrides", 2)
    bullet(doc, "Edit any extracted field on the review screen (number, dealer, amounts, line items).")
    bullet(doc, "Register a new dealer yourself if AI guessed the wrong supplier.")
    bullet(doc, "Skip AI entirely with Manual Invoice entry.")
    bullet(doc, "WhatsApp photos sit in an inbox until you tap Send to AI — then still need Verify.")
    bullet(doc, "confirm_matches checkbox is required — no silent auto-save.")

    h(doc, "6.2 Cheque stage overrides", 2)
    bullet(doc, "Compute bundles with pure Python buttons — no chat required.")
    bullet(doc, "Chat with the assistant — or ignore chat completely.")
    bullet(doc, "Manual bundling API/UI — assign invoices to groups yourself.")
    bullet(doc, "Reviewer suggestions — Apply only if you want; otherwise leave as-is.")
    bullet(doc, "Change dates, move invoices, split cheques via tools or forms.")
    bullet(doc, "Acknowledge warnings consciously if you still want to proceed.")
    bullet(doc, "Final Commit chooses bank account and cheque numbers — you own the print/commit step.")

    h(doc, "6.3 Other overrides", 2)
    bullet(doc, "Guide agent: only navigation; cannot change money decisions.")
    bullet(doc, "Analyst report: informational after the fact; does not alter committed cheques.")
    bullet(doc, "Cash Flow page: deterministic Python projections you can update with planned deposits.")
    bullet(doc, "Feature flags: USE_FAKE_AI, USE_BUNDLING_TOOL_AGENT=false (legacy JSON mode), etc.")

    callout(
        doc,
        "Bottom line for the team: If AI is wrong, the merchant can always edit the form, recompute manually, reject reviewer suggestions, or refuse to commit.",
    )

    # ------------------------------------------------------------------
    h(doc, "7. End-to-End Story (Example)", 1)
    numbered(doc, "Upload invoice photo → Vision fills a draft form.")
    numbered(doc, "You notice the total is wrong → you fix it → Verify.")
    numbered(doc, "On Cheques page, Python packs 5 invoices into 2 cheques under Rs. 500,000.")
    numbered(doc, "You chat: “Move invoice INV-3 to cheque 2.”")
    numbered(doc, "Assistant calls move_invoice(dry_run=true) → guardrails OK → preview shown.")
    numbered(doc, "You like it → apply / confirm → draft updates in session.")
    numbered(doc, "Reviewer suggests a later date → you ignore Apply.")
    numbered(doc, "Preview → Commit with cheque numbers → DB updated → Analyst report generated.")

    # ------------------------------------------------------------------
    h(doc, "8. Key Files (for developers)", 1)
    table(
        doc,
        ["Topic", "Where to look"],
        [
            ("Vision extract", "agents/ingestion.py"),
            ("Anomaly rules", "agents/anomaly.py"),
            ("Tool-calling chat", "agents/bundling_assistant.py + agents/bundling_tools.py"),
            ("Legacy JSON chat", "agents/assistant.py"),
            ("Reviewer", "agents/reviewer.py"),
            ("Guide", "agents/guide.py"),
            ("Analyst", "agents/analyst.py"),
            ("Guardrails", "core/guardrails.py"),
            ("Packing maths", "core/bundling.py, core/cheque_batcher.py, core/liquidity_engine.py"),
            ("HTTP wiring", "routes/ingestion.py, routes/bundling.py, routes/guide.py"),
            ("Optional pipeline", "agentic/orchestrator/pipeline.py"),
            ("Config flags", "config.py / .env"),
        ],
    )

    # ------------------------------------------------------------------
    h(doc, "9. Takeaways for Teammates", 1)
    bullet(doc, "Agents are helpers, not authorities.")
    bullet(doc, "Bundling Assistant uses real tool calling so maths stays in Python.")
    bullet(doc, "Guardrails catch ceiling breaches, bad dates, duplicate invoices, day-limit risk.")
    bullet(doc, "Drafts live in session until Verify / Commit.")
    bullet(doc, "Users can edit forms, manual-bundle, reject reviewer apply, and acknowledge warnings intentionally.")
    bullet(doc, "When explaining Zenith to others: “AI suggests → Python validates → human decides.”")

    footer = doc.add_paragraph()
    footer.add_run(
        "Document generated from the Zenith codebase for internal teammate onboarding. "
        "See also README.md and database/DATABASE.md."
    ).italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
